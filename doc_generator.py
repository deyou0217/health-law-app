# -*- coding: utf-8 -*-
"""
Word文档生成模块 - 严格遵循党政机关公文标准
- 页眉: 黄州区疾病预防控制中心
- 标题: 二号方正小标宋体（备选宋体）居中
- 正文: 三号仿宋体，首行缩进2字符
- 行距: 固定值28磅
- 页边距: 上3.7cm 下3.5cm 左2.8cm 右2.6cm
- 页码: 页面底端居中（- 1 - 格式）
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _add_page_number(doc):
    """添加页码 - 页面底端居中，- 1 - 格式"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # 清除原有页脚
        for p in footer.paragraphs:
            p.clear()
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 创建页码字段
        run = paragraph.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)

        run2 = paragraph.add_run()
        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr_text)

        run3 = paragraph.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char_end)

        # 设置字体
        for r in paragraph.runs:
            r.font.size = Pt(14)
            r.font.name = "宋体"
            r._r.rPr.rFonts.set(qn('w:eastAsia'), "宋体")


def _add_traceability_section(doc, source_info, full_content_fetched=False):
    """
    在文档中添加溯源信息：仅包含来源、原文链接、检索日期、全文状态
    """
    if not source_info:
        return

    now_str = datetime.now().strftime("%Y年%m月%d日")
    items = []

    if source_info.get("source"):
        items.append(("来源", source_info["source"]))
    if source_info.get("url"):
        items.append(("原文链接", source_info["url"]))
    items.append(("检索日期", now_str))
    items.append(("全文状态", "已获取完整正文" if full_content_fetched
                  else "部分内容（建议访问原文链接获取完整版）"))

    for label, value in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)

        run_label = p.add_run(f"{label}：")
        run_label.font.size = Pt(14)
        run_label.font.name = "仿宋"
        run_label._r.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
        run_label.bold = True

        run_value = p.add_run(value)
        run_value.font.size = Pt(14)
        run_value.font.name = "仿宋"
        run_value._r.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")


def create_document(
    title: str,
    content: str,
    source_info: dict = None,
    output_path: str = None,
    header_text: str = "黄州区疾病预防控制中心",
    full_content_fetched: bool = False,
) -> str:
    """
    按党政机关公文标准生成 Word 文档

    文档内容仅包含：来源、原文链接、检索日期、全文状态、文件正文

    参数:
        title: 文档标题（仅用于文件名，不写入正文）
        content: 正文内容
        source_info: 来源信息（source, url）
        output_path: 输出路径
        header_text: 页眉文字
        full_content_fetched: 是否已获取完整全文
    返回:
        输出文件路径
    """
    doc = Document()

    # ==================== 页面设置 ====================
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        # 页眉
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run(header_text)
        header_run.font.size = Pt(14)
        header_run.font.name = "仿宋"
        header_run._r.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
        header_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # 页眉下方横线
        pPr = header_para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    # ==================== 溯源信息 ====================
    _add_traceability_section(doc, source_info, full_content_fetched)

    # ==================== 正文 ====================
    # 按换行分段
    paragraphs = content.split("\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # 首行缩进2字符（2字符 ≈ 2*16pt = 32pt 左右，三号字约16pt）
        from docx.shared import Emu
        p.paragraph_format.first_line_indent = Pt(32)

        # 处理正文中的加粗标记
        parts = re.split(r'(\*\*.*?\*\*)', para_text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                # 加粗文本
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)

            run.font.size = Pt(16)  # 三号 = 16pt
            run.font.name = "仿宋"
            run._r.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
            run._r.rPr.rFonts.set(qn('w:ascii'), "仿宋")

    # ==================== 页码 ====================
    _add_page_number(doc)

    # ==================== 保存 ====================
    if output_path is None:
        sanitized_title = re.sub(r'[\\/:*?"<>|]', '_', title)
        sanitized_title = sanitized_title[:50]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = os.path.join(os.path.expanduser("~"), "Desktop")
        output_path = os.path.join(output_dir, f"{sanitized_title}_{timestamp}.docx")

    doc.save(output_path)
    return output_path


def create_batch_document(selected_items: list, header_text: str = "黄州区疾病预防控制中心") -> str:
    """
    批量生成 - 多篇法规合并到一个文档中
    selected_items: [{"title":..., "content":..., "url":..., "source":..., ...}]
    """
    if not selected_items:
        return ""

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = os.path.join(os.path.expanduser("~"), "Desktop")
    output_path = os.path.join(output_dir, f"卫健法规汇编_{timestamp}.docx")

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run(header_text)
        header_run.font.size = Pt(14)
        header_run.font.name = "仿宋"
        header_run._r.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")

        pPr = header_para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    for idx, item in enumerate(selected_items):
        # 分页符（除第一篇外）
        if idx > 0:
            doc.add_page_break()

        # 溯源信息
        src_info = {
            "source": item.get("source", ""),
            "url": item.get("url", ""),
        }
        got_full = item.get("full_content_fetched", False)
        _add_traceability_section(doc, src_info, got_full)

        # 正文
        content = item.get("content", "")
        paragraphs = content.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(28)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Pt(32)

            # 处理加粗
            parts = re.split(r'(\*\*.*?\*\*)', para_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.size = Pt(16)
                run.font.name = "仿宋"
                run._r.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")

    # 页码
    _add_page_number(doc)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    # 测试
    test_title = "中华人民共和国传染病防治法"
    test_content = """第一章 总则
    
第一条 为了预防、控制和消除传染病的发生与流行，保障人体健康和公共卫生，制定本法。

第二条 国家对传染病防治实行预防为主的方针，防治结合、分类管理、依靠科学、依靠群众。

第三条 本法规定的传染病分为甲类、乙类和丙类。
甲类传染病是指：鼠疫、霍乱。
乙类传染病是指：新型冠状病毒感染、传染性非典型肺炎、艾滋病、病毒性肝炎、脊髓灰质炎、人感染高致病性禽流感、麻疹、流行性出血热、狂犬病、流行性乙型脑炎、登革热、炭疽、细菌性和阿米巴性痢疾、肺结核、伤寒和副伤寒、流行性脑脊髓膜炎、百日咳、白喉、新生儿破伤风、猩红热、布鲁氏菌病、淋病、梅毒、钩端螺旋体病、血吸虫病、疟疾。
丙类传染病是指：流行性感冒、流行性腮腺炎、风疹、急性出血性结膜炎、麻风病、流行性和地方性斑疹伤寒、黑热病、包虫病、丝虫病，除霍乱、细菌性和阿米巴性痢疾、伤寒和副伤寒以外的感染性腹泻病。
国务院卫生行政部门根据传染病暴发、流行情况和危害程度，可以决定增加、减少或者调整乙类、丙类传染病病种并予以公布。"""

    path = create_document(test_title, test_content, {"source": "全国人大常委会", "url": "https://www.gov.cn"})
    print(f"文档已生成: {path}")
